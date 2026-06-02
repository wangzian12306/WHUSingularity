"""
分析 "安装维护手册 - 模板.docx" 的结构
使用 python-docx 读取并提取文档结构信息
"""
import sys
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DOC_PATH = "/Users/lubover/projects/whu-singularity/docs/安装维护手册 - 模板.docx"
OUTPUT_PATH = "/Users/lubover/projects/whu-singularity/docs/deploy/template_analysis.txt"


def analyze_document(path):
    doc = Document(path)

    lines = []
    lines.append("=" * 80)
    lines.append("文档分析: 安装维护手册 - 模板.docx")
    lines.append("=" * 80)

    # 1. 文档属性
    lines.append("\n--- 文档属性 ---")
    props = doc.core_properties
    lines.append(f"  标题: {props.title}")
    lines.append(f"  作者: {props.author}")
    lines.append(f"  创建时间: {props.created}")
    lines.append(f"  修改时间: {props.modified}")
    lines.append(f"  修订次数: {props.revision}")
    lines.append(f"  类别: {props.category}")

    # 2. 段落结构
    lines.append("\n--- 段落结构 (含样式) ---")
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "None"
        text = para.text[:120] if para.text else "(空)"
        # 判断是否为标题
        is_heading = style_name.startswith("Heading") or style_name.startswith("heading")
        prefix = "  " if not is_heading else ">> "
        lines.append(f"{prefix}[P{i}] style='{style_name}' | {text}")

    # 3. 表格结构
    lines.append("\n--- 表格结构 ---")
    for ti, table in enumerate(doc.tables):
        lines.append(f"\n[表格 {ti}] 行数={len(table.rows)}, 列数={len(table.columns)}")
        for ri, row in enumerate(table.rows):
            cells_text = []
            for ci, cell in enumerate(row.cells):
                ct = cell.text.replace("\n", "\\n")[:60]
                cells_text.append(f"  C{ci}='{ct}'")
            lines.append(f"  行{ri}: {';'.join(cells_text[:4])}")  # 最多显示4列

    # 4. 节信息
    lines.append("\n--- 节信息 ---")
    for si, section in enumerate(doc.sections):
        lines.append(f"  节{si}: 页面宽度={section.page_width}, 页面高度={section.page_height}")
        lines.append(f"    左页边距={section.left_margin}, 右页边距={section.right_margin}")
        lines.append(f"    页眉距离={section.header_distance}, 页脚距离={section.footer_distance}")
        # 页眉
        if section.header and section.header.paragraphs:
            header_text = " | ".join(p.text[:80] for p in section.header.paragraphs if p.text)
            lines.append(f"    页眉: {header_text}")
        # 页脚
        if section.footer and section.footer.paragraphs:
            footer_text = " | ".join(p.text[:80] for p in section.footer.paragraphs if p.text)
            lines.append(f"    页脚: {footer_text}")

    # 5. 样式列表
    lines.append("\n--- 文档中使用的样式 ---")
    used_styles = set()
    for para in doc.paragraphs:
        if para.style:
            used_styles.add(para.style.name)
    for sname in sorted(used_styles):
        lines.append(f"  {sname}")

    # 6. 图片/媒体引用
    lines.append("\n--- 媒体/图片引用 ---")
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
            lines.append(f"  图片: {rel.target_ref}")
    lines.append(f"  共 {image_count} 张图片")

    # 写入文件
    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print(f"分析完成，输出到: {OUTPUT_PATH}")
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")
    print(f"图片数: {image_count}")
    print(f"节数: {len(doc.sections)}")

    # 同时打印到 stdout
    print("\n".join(lines))


if __name__ == "__main__":
    analyze_document(DOC_PATH)
