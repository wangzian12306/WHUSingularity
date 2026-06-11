"""深入分析模板中的字体、段落格式等样式细节"""
from docx import Document
from docx.shared import Pt, Inches, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

doc = Document('docs/deploy/安装维护手册 - 模板.docx')

OUT = "docs/deploy/template_font_details.txt"
lines = []

# ---- 1. 样式定义（从 style 元素读取，而非 run） ----
lines.append("=" * 80)
lines.append("1. 样式定义 (style-level properties)")
lines.append("=" * 80)

# 从 paragraph 的 style 反向查样式的 XML
known_styles = set()
for para in doc.paragraphs:
    if para.style:
        known_styles.add(para.style.name)

for sname in sorted(known_styles):
    try:
        style = doc.styles[sname]
    except KeyError:
        continue
    lines.append(f"\n--- 样式: '{sname}' (type={style.type}) ---")

    # Paragraph format
    pf = style.paragraph_format
    lines.append(f"  paragraph_format.alignment = {pf.alignment}")
    lines.append(f"  paragraph_format.space_before = {pf.space_before}")
    lines.append(f"  paragraph_format.space_after = {pf.space_after}")
    lines.append(f"  paragraph_format.line_spacing = {pf.line_spacing}")
    lines.append(f"  paragraph_format.first_line_indent = {pf.first_line_indent}")
    lines.append(f"  paragraph_format.left_indent = {pf.left_indent}")

    # Font
    f = style.font
    lines.append(f"  font.name = {f.name}")
    lines.append(f"  font.size = {f.size}")
    lines.append(f"  font.bold = {f.bold}")
    lines.append(f"  font.italic = {f.italic}")
    lines.append(f"  font.underline = {f.underline}")
    if f.color and f.color.rgb:
        lines.append(f"  font.color.rgb = {f.color.rgb}")
    if f.color and f.color.theme_color:
        lines.append(f"  font.color.theme_color = {f.color.theme_color}")

    # East Asian font
    rPr = style.element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if rPr is not None:
        rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        if rFonts is not None:
            lines.append(f"  rFonts.eastAsia = {rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')}")
            lines.append(f"  rFonts.ascii = {rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')}")
            lines.append(f"  rFonts.hAnsi = {rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi')}")

# ---- 2. 表格样式详情 ----
lines.append("\n\n" + "=" * 80)
lines.append("2. 表格结构及内容 (用于对照填充)")
lines.append("=" * 80)

for ti, table in enumerate(doc.tables):
    lines.append(f"\n[表格 {ti}] {len(table.rows)}行 × {len(table.columns)}列")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            # 检查 cell 中段落样式
            cell_info = []
            for pi, para in enumerate(cell.paragraphs):
                style_name = para.style.name if para.style else 'None'
                text = para.text.replace('\n', '\\n')
                align = para.paragraph_format.alignment
                cell_info.append(f"P{pi}({style_name},align={align})='{text[:80]}'")
                # 字体
                for run in para.runs:
                    f = run.font
                    cell_info.append(f"  run: font={f.name} size={f.size} bold={f.bold}")
            lines.append(f"  行{ri}列{ci}: {'; '.join(cell_info)}")

# ---- 3. 每个段落的 run 级字体覆盖 ----
lines.append("\n\n" + "=" * 80)
lines.append("3. 段落 Run 级覆盖 (与样式不同的地方)")
lines.append("=" * 80)

for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else 'None'
    for run in para.runs:
        f = run.font
        overrides = []
        if f.name: overrides.append(f"font.name={f.name}")
        if f.size: overrides.append(f"size={f.size} ({f.size/12700}pt)")
        if f.bold is not None: overrides.append(f"bold={f.bold}")
        if f.italic is not None: overrides.append(f"italic={f.italic}")
        if f.color and f.color.rgb: overrides.append(f"color={f.color.rgb}")
        if overrides:
            lines.append(f"  [P{i}] style={style_name} | text='{run.text[:60]}' | {' '.join(overrides)}")

# ---- 4. 表格样式 ----
lines.append("\n\n" + "=" * 80)
lines.append("4. 表格使用的样式")
lines.append("=" * 80)
for ti, table in enumerate(doc.tables):
    tbl_style = table.style.name if table.style else 'None'
    lines.append(f"  Table {ti} style = {tbl_style}")

with open(OUT, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

print(f"Written to {OUT}")
print("\n".join(lines))
