"""
填充脚本：将 content_to_fill.md 的内容填入 安装维护手册 - 模板.docx
严格保留模板字体/样式，仅替换文本内容。
对超出模板段落数的章节，在当前位置之后插入新段落（克隆样式）。
"""
import copy
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

TEMPLATE = "docs/deploy/安装维护手册 - 模板.docx"
OUTPUT = "docs/deploy/WHUSingularity安装维护手册.docx"

doc = Document(TEMPLATE)

# ============================================================
# 辅助函数
# ============================================================

def set_run_text(run, text):
    """替换 run 文本，保留字体格式"""
    run.text = text

def set_paragraph_text(para, text, font_name="Times New Roman"):
    """设置段落全部文本（清空原有 runs，新建一个 run 并带 Times New Roman）"""
    # 删除所有 run
    for r in para.runs:
        r._element.getparent().remove(r._element)
    # 添加新 run
    run = para.add_run(text)
    run.font.name = font_name
    # 设置东亚字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    return run

def set_paragraph_with_bold_steps(para, text: str):
    """
    设置段落文本，处理 '第X步：' 模式：数字部分加粗，其余正常。
    text 包含多步时用换行分隔，如: '第1步：xxx\\n第2步：yyy'
    """
    # 清空原有
    for r in para.runs:
        r._element.getparent().remove(r._element)

    import re
    segments = re.split(r'(第\d+步)', text)
    for seg in segments:
        if not seg:
            continue
        if re.match(r'第\d+步', seg):
            run = para.add_run(seg + "：")
            run.font.name = "Times New Roman"
            run.bold = True
            # 设置东亚字体
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), "Times New Roman")
            rFonts.set(qn('w:ascii'), "Times New Roman")
            rFonts.set(qn('w:hAnsi'), "Times New Roman")
            rPr.insert(0, rFonts)
        elif seg.startswith("："):
            pass  # skip standalone colon, already added above
        else:
            # Normal text after step description
            run = para.add_run(seg)
            run.font.name = "Times New Roman"
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), "Times New Roman")
            rFonts.set(qn('w:ascii'), "Times New Roman")
            rFonts.set(qn('w:hAnsi'), "Times New Roman")
            rPr.insert(0, rFonts)

def add_run_to_paragraph(para, text, bold=False, font_name="Times New Roman"):
    """向段落添加一个 run"""
    run = para.add_run(text)
    run.font.name = font_name
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)
    return run

def clone_paragraph_after(para):
    """
    在 para 之后插入一个同样式的新段落，返回新段落。
    通过复制 XML 元素实现。
    """
    new_p = OxmlElement('w:p')
    # 复制段落属性（样式引用等）
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        new_pPr = copy.deepcopy(pPr)
        # 移除编号属性避免重复编号
        numPr = new_pPr.find(qn('w:numPr'))
        if numPr is not None:
            new_pPr.remove(numPr)
        new_p.append(new_pPr)
    
    # 插入到 para 之后
    para._element.addnext(new_p)
    
    # 创建 python-docx Paragraph 对象
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, para._parent)
    # 更新 internal _element 引用（Paragraph 包装的是 CT_P，需要关联正确的 XML 元素）
    # 实际返回 Paragraph 对象即可
    return new_para

def insert_paragraphs_after(anchor_para, style_name, texts):
    """
    在 anchor_para 之后插入多个新段落（指定样式），
    返回新段落列表（Paragraph 对象）。
    texts: list of (style_name, text_content)
    """
    from docx.text.paragraph import Paragraph
    result = []
    current_element = anchor_para._element
    parent = anchor_para._parent
    
    # 样式名 → 样式ID 映射
    style_id_map = {
        'Heading 1': '1', 'Heading 2': '2', 'Heading 3': '3',
        'Normal': 'Normal', 'List Paragraph': 'ListParagraph',
    }
    for (sty_name, text) in texts:
        new_p = OxmlElement('w:p')
        # 设置样式
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_id_map.get(sty_name, sty_name))
        pPr.append(pStyle)
        new_p.append(pPr)
        
        # 插入到当前元素之后
        current_element.addnext(new_p)
        current_element = new_p  # 下次插入在这个之后
        
        new_para = Paragraph(new_p, parent)
        
        # 添加文本
        set_paragraph_text(new_para, text)
        result.append(new_para)
    
    return result

def is_merge_continuation(cell):
    """检查单元格是否为纵向合并的延续（vMerge='continue'），这类单元格不应覆写"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        vMerge = tcPr.find(qn('w:vMerge'))
        if vMerge is not None and vMerge.get(qn('w:val')) == 'continue':
            return True
    return False

def replace_cell_text(cell, text):
    """
    替换单元格全部文本。通过修改现有 run 文本而非删除重建，
    避免合并单元格的 DOM 交叉干扰。
    若单元格为空（无 run），则新建一个 run。
    """
    if is_merge_continuation(cell):
        return
    all_runs = []
    for para in cell.paragraphs:
        all_runs.extend(para.runs)
    
    if all_runs:
        # 第一个 run 设为目标文本
        all_runs[0].text = text
        # 清空其余 run
        for run in all_runs[1:]:
            run.text = ""
    else:
        # 无 run，添加一个
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), "Times New Roman")
        rFonts.set(qn('w:ascii'), "Times New Roman")
        rFonts.set(qn('w:hAnsi'), "Times New Roman")
        rPr.insert(0, rFonts)

def replace_table_cells(table, cell_map):
    """按 (row, col) → text 映射替换表格内容。"""
    for (ri, ci), text in cell_map.items():
        cell = table.rows[ri].cells[ci]
        replace_cell_text(cell, text)

# ============================================================
# PHASE 1: 段落文本替换
# ============================================================

paras = doc.paragraphs

# --- 封面 ---
set_paragraph_text(paras[0], "WHUSingularity")   # P0: XXX项目
# P1: 安装维护手册 - keep
set_paragraph_text(paras[1], "安装维护手册")

# --- 变更履历区域 (P4-P8) ---
# P4 "变更履历" - 保留
# P5 "说明："变更原因"主要是分为：" - 保留
# P6-P8 List Paragraph - 保留

# --- 1. 引言 ---
set_paragraph_text(paras[11], "引言")  # Heading 1
set_paragraph_text(paras[12], "编写目的")  # Heading 2

purpose_text = (
    "本文档旨在为 WHUSingularity 高并发秒杀系统的部署与运维人员提供完整的安装、配置和维护指南。"
    "通过本文档，读者可以："
    "（1）了解系统的整体架构和组件依赖关系；"
    "（2）掌握基础设施环境（MySQL、Redis、Nacos、RocketMQ）的搭建方法；"
    "（3）完成各微服务模块的编译、配置与部署；"
    "（4）掌握日常运维操作，包括启停、监控、日志查看、故障排查与扩容缩容。"
    "预期读者包括系统运维工程师、后端开发工程师及项目管理人员。"
)
set_paragraph_text(paras[13], purpose_text)

set_paragraph_text(paras[14], "软件背景")

bg_text = (
    "WHUSingularity 是一套基于 Spring Cloud 微服务架构的高并发抢单（秒杀）系统。"
    "系统核心为 singularity-core 自定义的高并发资源分配框架，"
    "上层业务为电商秒杀场景，配套 React + TypeScript 前端。"
    "技术栈：Java 21、Spring Boot 3.2.6、Spring Cloud 2023.0.3、"
    "Spring Cloud Alibaba 2023.0.3.2、Nacos（服务注册与配置中心）、"
    "Spring Cloud Gateway（API 网关）、OpenFeign（服务间调用）、"
    "MyBatis 3.0.4、MySQL 8.0、Redis 7.x、RocketMQ 5.x、"
    "Caffeine（本地缓存）、Flyway（数据库迁移）、React 19 + Vite + Ant Design（前端）、"
    "Docker + Docker Compose（容器化部署）。"
)
set_paragraph_text(paras[15], bg_text)

# --- 2. 环境准备 ---
set_paragraph_text(paras[16], "环境准备")  # Heading1: "操作系统安装" → "环境准备"
set_paragraph_text(paras[17], "硬件环境")

# P18 was empty → fill hardware text
hw_text = (
    "推荐配置（开发/测试环境）：CPU 4 核及以上，内存 16 GB 及以上，"
    "磁盘 50 GB 可用空间，操作系统 Linux（Ubuntu 22.04 / CentOS 7+）或 macOS 12+，"
    "各服务端口需互通。生产环境建议至少 3 台节点，每节点 8 核 / 32 GB，"
    "配合容器编排（Kubernetes / Docker Swarm）。"
)
set_paragraph_text(paras[18], hw_text)

# P19: "Linux系统安装" → "基础软件环境"
set_paragraph_text(paras[19], "基础软件环境")

# P20: "简单简介" → base software intro
sw_intro = (
    "部署 WHUSingularity 前需安装以下基础软件：Docker 24.0+（容器运行时），"
    "Docker Compose v2.20+（多容器编排），JDK 21 OpenJDK（Java 运行时），"
    "Maven 3.9+（项目构建），Git 2.40+（源码管理）。"
)
set_paragraph_text(paras[20], sw_intro)

# P21: "准备工作" → "Docker 安装"
set_paragraph_text(paras[21], "Docker 安装")

# P22: "简单说明" → Docker install intro
docker_intro = (
    "Docker 用于统一运行基础设施组件（MySQL、Redis、Nacos、RocketMQ）及各微服务容器。"
    "以下以 Ubuntu 22.04 为例，macOS 可直接安装 Docker Desktop。"
)
set_paragraph_text(paras[22], docker_intro)

# P23: "安装过程" → "JDK 21 安装" (changed hierarchy)
# Actually, the template has:
#   P21: Heading3 "准备工作"
#   P22: Normal "简单说明"
#   P23: Heading3 "安装过程"
#   P24: Normal "按步骤介绍安装过程并附安装截图"
#   P25: Normal "第1步："
#   P26: Normal "第2步："
#   P27: Normal "第3步："
# Then P28: Heading1 "Oracle数据库安装"
#
# For Docker install, I want to use P23-P27 for Docker steps (need 7 steps but only 3 available).
# I'll put JDK in new paragraphs after Docker, and Maven after JDK.

# P23: "安装过程" → keep as general heading or change to "安装步骤"
set_paragraph_text(paras[23], "安装步骤")

# P24: combined Docker step descriptions
# Clear P24 and add as intro
set_paragraph_text(paras[24], "以下为 Docker 在 Ubuntu 22.04 上的安装步骤：")

# P25-P27: Docker steps (we have 7, but only 3 paragraphs)
# Strategy: put steps 1-3 in P25-P27, then insert P27 clones for steps 4-7
docker_steps = [
    "第1步：卸载旧版本。执行 sudo apt-get remove docker docker-engine docker.io containerd runc",
    "第2步：安装依赖。执行 sudo apt-get update && sudo apt-get install -y ca-certificates curl gnupg lsb-release",
    "第3步：添加 Docker 官方 GPG 密钥。sudo mkdir -p /etc/apt/keyrings && curl -fsSL https://download.docker.com/linux/ubuntu/gpg | sudo gpg --dearmor -o /etc/apt/keyrings/docker.gpg",
    "第4步：添加 Docker 仓库。echo \"deb [arch=$(dpkg --print-architecture) signed-by=/etc/apt/keyrings/docker.gpg] https://download.docker.com/linux/ubuntu $(lsb_release -cs) stable\" | sudo tee /etc/apt/sources.list.d/docker.list > /dev/null",
    "第5步：安装 Docker Engine。sudo apt-get update && sudo apt-get install -y docker-ce docker-ce-cli containerd.io docker-compose-plugin",
    "第6步：验证安装。执行 sudo docker --version && sudo docker compose version",
    "第7步（可选）：将当前用户加入 docker 组，免 sudo 执行。sudo usermod -aG docker $USER && newgrp docker",
]

set_paragraph_text(paras[25], docker_steps[0])
set_paragraph_text(paras[26], docker_steps[1])
set_paragraph_text(paras[27], docker_steps[2])

# Insert steps 4-7 after P27
last_step_para = paras[27]
for step in docker_steps[3:]:
    new_p = clone_paragraph_after(last_step_para)
    set_paragraph_text(new_p, step)
    last_step_para = new_p

# Now insert JDK 21 section
# Insert after the last Docker step paragraph
jdk_sections = [
    ("Heading 3", "JDK 21 安装"),
    ("Normal", "安装 OpenJDK 21 作为 Java 运行时环境。"),
    ("Normal", "第1步：安装 OpenJDK 21。执行 sudo apt-get update && sudo apt-get install -y openjdk-21-jdk"),
    ("Normal", "第2步：验证安装。执行 java -version，预期输出 openjdk version \"21.0.x\""),
    ("Normal", "第3步：配置 JAVA_HOME（可选）。echo 'export JAVA_HOME=/usr/lib/jvm/java-21-openjdk-amd64' >> ~/.bashrc && echo 'export PATH=\\$JAVA_HOME/bin:\\$PATH' >> ~/.bashrc && source ~/.bashrc"),
]

new_paras_jdk = insert_paragraphs_after(last_step_para, "Normal", jdk_sections)
last_step_para = new_paras_jdk[-1]

# Insert Maven section
maven_sections = [
    ("Heading 3", "Maven 安装"),
    ("Normal", "Maven 用于项目构建与依赖管理。"),
    ("Normal", "第1步：下载 Maven 3.9.9。wget https://dlcdn.apache.org/maven/maven-3/3.9.9/binaries/apache-maven-3.9.9-bin.tar.gz"),
    ("Normal", "第2步：解压至 /opt。sudo tar -xzf apache-maven-3.9.9-bin.tar.gz -C /opt && sudo ln -s /opt/apache-maven-3.9.9 /opt/maven"),
    ("Normal", "第3步：配置环境变量。echo 'export M2_HOME=/opt/maven' >> ~/.bashrc && echo 'export PATH=\\$M2_HOME/bin:\\$PATH' >> ~/.bashrc && source ~/.bashrc"),
    ("Normal", "第4步：验证安装。mvn --version"),
]

new_paras_maven = insert_paragraphs_after(last_step_para, "Normal", maven_sections)
last_step_para = new_paras_maven[-1]

# --- 3. 基础设施安装 (was "Oracle数据库安装") ---
# P28: Heading1 "Oracle数据库安装" → "基础设施安装"
# P29: Heading2 "前期准备" → keep concept
# P30: Normal "简单说明"
# P31: Heading2 "安装流程"
# P32: Normal "按步骤..."
# P33-P35: steps

set_paragraph_text(paras[28], "基础设施安装")

prep_text = (
    "WHUSingularity 依赖以下基础设施组件，全部通过 Docker Compose 一键部署："
    "MySQL 8.0（端口 3306，持久化存储）、Redis 7.x（端口 6379，缓存与分布式计数）、"
    "Nacos 2.3.x（端口 8848/9848，服务注册发现与配置中心）、"
    "RocketMQ NameServer 5.x（端口 9876）、RocketMQ Broker 5.x（端口 10911/10909）。"
)
set_paragraph_text(paras[29], "基础设施组件")
set_paragraph_text(paras[30], prep_text)

set_paragraph_text(paras[31], "安装流程")

infra_steps = [
    "第1步：获取部署文件。项目源码中 deploy/ 目录包含 Docker Compose 编排文件 docker-compose-infra.yml。进入目录执行 ls docker-compose*.yml 确认。",
    "第2步：创建持久化数据目录。执行 mkdir -p /data/mysql /data/redis /data/nacos /data/rocketmq，并确保以下端口未被占用：3306、6379、8848、9848、9876、10911、10909。",
    "第3步：启动基础设施。执行 docker compose -f docker-compose-infra.yml up -d，等待所有容器启动完成（约 30-60 秒）。",
    "第4步：验证各组件状态。MySQL 验证：docker exec mysql-container mysqladmin ping -h localhost -u root -proot123；Redis 验证：docker exec redis-container redis-cli ping（应返回 PONG）；Nacos 验证：curl http://localhost:8848/nacos/v1/console/health/readiness（应返回 ok）。",
    "第5步：初始化 Nacos 配置。登录 Nacos 控制台（http://localhost:8848/nacos，默认账号/密码：nacos/nacos），创建 singularity-gateway.yaml、singularity-user.yaml、singularity-order.yaml、singularity-stock.yaml、singularity-product.yaml、singularity-scaler.yaml 等配置文件。",
    "第6步：初始化数据库。执行 docker exec mysql-container mysql -u root -proot123 创建 whu_user、whu_order、whu_stock、whu_product 数据库。数据库迁移由各服务启动时通过 Flyway 自动执行。",
]
set_paragraph_text(paras[32], "以下为通过 Docker Compose 部署基础设施的详细步骤：")
set_paragraph_text(paras[33], infra_steps[0])
set_paragraph_text(paras[34], infra_steps[1])
set_paragraph_text(paras[35], infra_steps[2])

# Insert steps 4-6 after P35
last_infra = paras[35]
for step in infra_steps[3:]:
    new_p = clone_paragraph_after(last_infra)
    set_paragraph_text(new_p, step)
    last_infra = new_p

# --- 4. 项目应用部署 ---
# P36: Heading1 "项目应用部署" → keep
# P37: Heading2 "WEB应用程序发布" → "应用程序构建"
# P38: Normal "简要说明"
# P39: Heading3 "JDK的安装" → "获取源码"
# P40: Normal step intro
# P41-43: steps
# P44: Heading3 "Tomcat的安装" → "编译打包"
# P45-48: steps
# P49: Heading3 "Web应用程序部署" → "前端构建"
# P50-53: steps
# P54: Heading2 "数据库部署" → keep
# P55-58: steps
# P59: Heading2 "系统部署信息汇总" → keep
# P60: Normal bold "系统部署信息汇总表"
# P61: Normal empty
# P62: Heading1 "系统管理员平台维护" → "系统运维管理"
# P63-P75: maintenance sections

set_paragraph_text(paras[36], "项目应用部署")

# 4.1 应用程序构建
set_paragraph_text(paras[37], "应用程序构建")
set_paragraph_text(paras[38], "使用 Maven 编译打包各微服务模块，使用 npm 构建前端。")

# 4.1.1 获取源码
set_paragraph_text(paras[39], "获取源码")
set_paragraph_text(paras[40], "获取项目源码并编译打包：")
set_paragraph_text(paras[41], "第1步：克隆代码仓库。git clone <repository-url> /opt/whu-singularity && cd /opt/whu-singularity")
set_paragraph_text(paras[42], "第2步：确认 Java 和 Maven 版本。java -version（应为 21）&& mvn --version（应为 3.9+）")
set_paragraph_text(paras[43], "第3步：可选。检查项目结构，确认各模块 pom.xml 依赖完整。")

# 4.1.2 编译打包
set_paragraph_text(paras[44], "编译打包")
set_paragraph_text(paras[45], "使用 Maven 全量构建各模块：")
set_paragraph_text(paras[46], "第1步：跳过测试全量构建。执行 mvn clean package -DskipTests")
set_paragraph_text(paras[47], "第2步：验证构建产物。执行 ls singularity-user/target/*.jar、ls singularity-order/target/*.jar 等，确认各模块 Jar 包生成成功。")
set_paragraph_text(paras[48], "第3步：构建单个模块（可选）。例如仅构建订单服务：mvn -pl singularity-order package -DskipTests")

# 4.1.3 前端构建
set_paragraph_text(paras[49], "前端构建")
set_paragraph_text(paras[50], "构建 React 前端应用：")
set_paragraph_text(paras[51], "第1步：进入前端目录。cd singularity-front")
set_paragraph_text(paras[52], "第2步：安装依赖。npm install")
set_paragraph_text(paras[53], "第3步：生产构建。npm run build，产物输出至 singularity-front/dist/ 目录。")

# 4.2 数据库部署 → "Docker Compose 一键部署"
set_paragraph_text(paras[54], "Docker Compose 一键部署")
set_paragraph_text(paras[55], "项目根目录提供 dev-run.sh 和 deploy/docker-compose.yml 可用于一键拉起所有服务：")
set_paragraph_text(paras[56], "第1步：确认基础设施已就绪。docker compose -f deploy/docker-compose-infra.yml ps 确认所有组件运行正常。")
set_paragraph_text(paras[57], "第2步：启动全部微服务。执行 ./dev-run.sh 或 docker compose -f deploy/docker-compose.yml up -d。")
set_paragraph_text(paras[58], "第3步：验证所有服务。curl http://localhost:8080/actuator/health 检查网关健康状态；通过网关访问各服务：curl http://localhost:8080/api/user/actuator/health。")

# After P58, insert: 4.3 各服务单独部署（可选）+ 4.4 数据库部署 + 4.5 部署信息汇总
# Actually, the template already has P59 (Heading2 "系统部署信息汇总"), P60, P61
# But we also need "各服务单独部署" and "数据库部署" sections
# Let me insert "各服务单独部署" between P58 and P59

extra_deploy = [
    ("Heading 2", "各服务单独部署（可选）"),
    ("Normal", "若不使用 Docker，可直接以 Jar 包方式按以下顺序启动（Nacos 必须先启动）：Nacos → User(8090) → Stock(8082) → Order(8081) → Product(8087) → Merchant(8091) → Gateway(8080) → Scaler(9090)。各服务启动命令格式为：java -jar -Dserver.port=<port> <module>/target/<module>-*.jar &。例如 singularity-user：java -jar -Dserver.port=8090 singularity-user/target/singularity-user-*.jar &。"),
    ("Heading 2", "数据库部署"),
    ("Normal", "数据库使用 MySQL 8.0（已在第 3 章通过 Docker 部署），各服务通过 Flyway 自动管理表结构迁移。Flyway 迁移脚本位于 singularity-stock/src/main/resources/db/migration/ 和 singularity-product/src/main/resources/db/migration/。首次启动服务时 Flyway 自动执行迁移，迁移历史记录存储在数据库的 flyway_schema_history 表中。"),
]

insert_paragraphs_after(paras[58], "Normal", extra_deploy)

# P59: "系统部署信息汇总"
set_paragraph_text(paras[59], "系统部署信息汇总")
set_paragraph_text(paras[60], "系统部署信息汇总表")  # keep bold via run
# P61 empty - keep

# --- 5. 系统运维管理 ---
set_paragraph_text(paras[62], "系统运维管理")

maintenance_intro = (
    "本章介绍 WHUSingularity 系统的日常运维操作，包括 Nacos 配置管理、服务启停、"
    "日志查看、健康检查与监控、常见问题排查、备份与恢复、以及扩容缩容等内容。"
)
set_paragraph_text(paras[63], maintenance_intro)

# 5.1 Nacos 配置管理
set_paragraph_text(paras[64], "Nacos 配置管理")
nacos_mgmt = (
    "Nacos 控制台地址：http://<host>:8848/nacos（默认账号/密码：nacos/nacos）。"
    "常用配置文件包括：singularity-order.yaml（订单服务核心配置，含 product-id、Redis、RocketMQ）、"
    "singularity-user.yaml（JWT 密钥与 Redis 连接）、singularity-stock.yaml（MQ 消费者配置）、"
    "singularity-gateway.yaml（路由规则）、singularity-scaler.yaml（自动伸缩策略）。"
    "修改配置后在 Nacos 控制台发布，对应服务会自动刷新（热更新，无需重启）。"
)
set_paragraph_text(paras[65], nacos_mgmt)

# 5.2 服务启停
set_paragraph_text(paras[66], "服务启停")
startstop = (
    "全部服务（Docker Compose）：启动 docker compose -f deploy/docker-compose.yml up -d；"
    "停止 docker compose -f deploy/docker-compose.yml stop；"
    "重启 docker compose -f deploy/docker-compose.yml restart；"
    "停止并删除 docker compose -f deploy/docker-compose.yml down。"
    "单个服务：docker restart/stop/start <container-name>。"
)
set_paragraph_text(paras[67], startstop)

# 5.3 日志查看
set_paragraph_text(paras[68], "日志查看")
log_text = (
    "Docker 部署：查看所有服务日志 docker compose -f deploy/docker-compose.yml logs -f；"
    "查看单个服务日志 docker compose -f deploy/docker-compose.yml logs -f singularity-order；"
    "查看最近 100 行 docker compose logs --tail 100 singularity-order。"
    "Jar 包部署：日志默认输出到标准输出，可通过 java -jar singularity-order.jar > logs/order.log 2>&1 & 持久化。"
)
set_paragraph_text(paras[69], log_text)

# 5.4 健康检查与监控
set_paragraph_text(paras[70], "健康检查与监控")
health_text = (
    "网关聚合健康检查：curl http://localhost:8080/actuator/health；"
    "各服务独立检查：curl http://localhost:8081/actuator/health（order）、"
    "curl http://localhost:8082/actuator/health（stock）等。"
    "Prometheus 指标端点（供自动伸缩服务使用）：curl http://localhost:8081/actuator/prometheus。"
    "Nacos 服务列表：登录 Nacos 控制台 → 服务管理 → 服务列表，查看所有微服务健康状态。"
)
set_paragraph_text(paras[71], health_text)

# 5.5 常见问题排查
set_paragraph_text(paras[72], "常见问题排查")
trouble_text = (
    "常见问题及解决方案："
    "（1）服务启动失败——确认 Nacos 已在 8848 端口运行；"
    "（2）数据库连接失败——检查 docker ps 和配置中的数据库连接信息；"
    "（3）MQ 消息发送失败——确认 NameServer(9876) 和 Broker(10911) 端口可达；"
    "（4）网关返回 503——检查目标服务是否已注册到 Nacos；"
    "（5）Redis 连接失败——执行 redis-cli ping 验证；"
    "（6）库存扣减异常——检查 Nacos 配置中 product-id 是否正确；"
    "（7）Flyway 迁移失败——查看服务启动日志中的 Flyway 错误详情。"
)
set_paragraph_text(paras[73], trouble_text)

# 5.6 备份与恢复
set_paragraph_text(paras[74], "备份与恢复")
backup_text = (
    "MySQL 数据备份：docker exec mysql-container mysqldump -u root -proot123 --all-databases > backup_$(date +%Y%m%d).sql；"
    "MySQL 数据恢复：docker exec -i mysql-container mysql -u root -proot123 < backup.sql；"
    "Redis 数据备份：docker cp redis-container:/data/dump.rdb ./redis_backup.rdb；"
    "Nacos 配置备份：在 Nacos 控制台逐个导出配置文件，或备份 Nacos 数据目录。"
)
set_paragraph_text(paras[75], backup_text)

# Insert 5.7 扩容与缩容 after P75
extra_maintenance = [
    ("Heading 2", "扩容与缩容"),
    ("Normal", "singularity-scaler 服务提供基于 Prometheus 指标的自动伸缩能力。自动伸缩策略通过 Nacos singularity-scaler.yaml 配置：threshold-cpu（CPU阈值%）、threshold-qps（QPS阈值）、min-instances/max-instances（实例数范围）、scale-up-cooldown/scale-down-cooldown（冷却时间秒）、port-start（动态实例起始端口）。手动扩容示例：docker run -d --name singularity-order-2 -e SERVER_PORT=8181 singularity-order:latest。"),
]

insert_paragraphs_after(paras[75], "Normal", extra_maintenance)

# ============================================================
# PHASE 2: 表格替换（使用去重写入，自动处理合并单元格）
# ============================================================

tables = doc.tables

# 表格0：封面信息 (5行×3列)
# C2 列填入对应值
replace_cell_text(tables[0].rows[0].cells[2], "内部")
replace_cell_text(tables[0].rows[1].cells[2], "WHU-SING-DEPLOY-001")
replace_cell_text(tables[0].rows[2].cells[2], "研发部")
replace_cell_text(tables[0].rows[3].cells[2], "2026-05")
replace_cell_text(tables[0].rows[4].cells[2], "V1.0")

# 表格1：签字栏 - 保持不变（留空）

# 表格2：变更履历 (5行×8列) — 填充第一行数据
t2 = tables[2]
replace_cell_text(t2.rows[1].cells[0], "1")
replace_cell_text(t2.rows[1].cells[1], "2026-05-30")
replace_cell_text(t2.rows[1].cells[2], "V1.0")
replace_cell_text(t2.rows[1].cells[3], "全文")
replace_cell_text(t2.rows[1].cells[4], "建立初稿")
replace_cell_text(t2.rows[1].cells[5], "初始版本")
replace_cell_text(t2.rows[1].cells[6], "—")
replace_cell_text(t2.rows[1].cells[7], "—")

# 表格3：硬件环境 (5行×2列) — Row 0 两列已合并，只写 C0
t3 = tables[3]
replace_cell_text(t3.rows[0].cells[0], "硬件配置（开发/测试 & 生产）")
# Row 0 C1 自动跳过（与 C0 同一 tc）
replace_cell_text(t3.rows[1].cells[0], "品牌/型号")
replace_cell_text(t3.rows[1].cells[1], "普通服务器 / 云主机")
replace_cell_text(t3.rows[2].cells[0], "CPU")
replace_cell_text(t3.rows[2].cells[1], "8 核及以上")
replace_cell_text(t3.rows[3].cells[0], "内存")
replace_cell_text(t3.rows[3].cells[1], "32 GB 及以上")
replace_cell_text(t3.rows[4].cells[0], "磁盘")
replace_cell_text(t3.rows[4].cells[1], "SSD 100 GB 及以上")

# 表格4：部署信息汇总 (6行×3列)
# C0 列有纵向合并：每两行共享同一 tc (0-1: JDK, 2-3: MySQL, 4-5: Oracle→Nacos/Redis)
# 因 C0 合并导致写一次会影响两行，我们统一写成组合文本或只改非合并的 C1/C2
t4 = tables[4]
replace_cell_text(t4.rows[0].cells[0], "JDK")
replace_cell_text(t4.rows[0].cells[1], "版本")
replace_cell_text(t4.rows[0].cells[2], "21 (OpenJDK)")
replace_cell_text(t4.rows[1].cells[1], "安装路径")
replace_cell_text(t4.rows[1].cells[2], "/usr/lib/jvm/java-21-openjdk")

replace_cell_text(t4.rows[2].cells[0], "MySQL")
replace_cell_text(t4.rows[2].cells[1], "版本")
replace_cell_text(t4.rows[2].cells[2], "8.0 (Docker)")
replace_cell_text(t4.rows[3].cells[1], "端口/路径")
replace_cell_text(t4.rows[3].cells[2], "3306")

replace_cell_text(t4.rows[4].cells[0], "Nacos / Redis")
replace_cell_text(t4.rows[4].cells[1], "版本")
replace_cell_text(t4.rows[4].cells[2], "2.3.x (Docker)")
replace_cell_text(t4.rows[5].cells[1], "版本")
replace_cell_text(t4.rows[5].cells[2], "7.x (Docker)")

# ============================================================
# PHASE 3: 页眉页脚
# ============================================================

for section in doc.sections:
    # 页脚
    footer = section.footer
    if footer and footer.paragraphs:
        for para in footer.paragraphs:
            full_text = para.text
            if "XXX有限公司版权所有" in full_text or "XXX服务有限公司" in full_text:
                # Clear and rewrite
                for run in para.runs:
                    run._element.getparent().remove(run._element)
                new_text = full_text.replace("XXX有限公司", "WHUSingularity 项目组").replace("XXX服务有限公司", "WHUSingularity 项目组")
                add_run_to_paragraph(para, new_text, font_name="Times New Roman")
            elif "XXX有限公司" in full_text or "XXX Co.,Ltd" in full_text:
                for run in para.runs:
                    run._element.getparent().remove(run._element)
                new_text = full_text.replace("XXX有限公司", "WHUSingularity 项目组").replace("XXX Co.,Ltd", "高并发秒杀系统")
                add_run_to_paragraph(para, new_text, font_name="Times New Roman")
    
    # 页眉 - keep as is (empty/tab)

# ============================================================
# PHASE 4: 更新目录（SDT 中的 TOC）
# ============================================================
# 模板的 TOC 存在于 body 的 <w:sdt> 元素中（body 子元素 [13]）
# 替换为新的 TOC 域代码，Word 打开时会自动/手动更新

body = doc.element.body
nsmap_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# 找到包含 TOC 的 SDT 元素
sdt_to_replace = None
for child in body:
    if child.tag == f'{{{nsmap_w}}}sdt':
        # 检查是否包含 TOC
        sdt_content = child.find(f'{{{nsmap_w}}}sdtContent')
        if sdt_content is not None:
            instr = sdt_content.findall(f'.//{{{nsmap_w}}}instrText')
            for it in instr:
                if it.text and 'TOC' in it.text:
                    sdt_to_replace = child
                    break
    if sdt_to_replace is not None:
        break

if sdt_to_replace is not None:
    from docx.oxml import OxmlElement
    
    # 构建新的 sdtContent：目录标题 + TOC 域
    new_content = OxmlElement('w:sdtContent')
    
    # 1) "目录" 标题段落
    p_title = OxmlElement('w:p')
    pPr_t = OxmlElement('w:pPr')
    pStyle_t = OxmlElement('w:pStyle')
    pStyle_t.set(f'{{{nsmap_w}}}val', 'TOC')
    pPr_t.append(pStyle_t)
    jc_t = OxmlElement('w:jc')
    jc_t.set(f'{{{nsmap_w}}}val', 'center')
    pPr_t.append(jc_t)
    p_title.append(pPr_t)
    r_title = OxmlElement('w:r')
    rPr_title = OxmlElement('w:rPr')
    rFonts_title = OxmlElement('w:rFonts')
    rFonts_title.set(f'{{{nsmap_w}}}ascii', 'Times New Roman')
    rFonts_title.set(f'{{{nsmap_w}}}hAnsi', 'Times New Roman')
    rFonts_title.set(f'{{{nsmap_w}}}cs', 'Times New Roman')
    rPr_title.append(rFonts_title)
    r_title.append(rPr_title)
    t_title = OxmlElement('w:t')
    t_title.text = '目录'
    t_title.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_title.append(t_title)
    p_title.append(r_title)
    new_content.append(p_title)
    
    # 2) TOC 域段落
    p_toc = OxmlElement('w:p')
    pPr_toc = OxmlElement('w:pPr')
    pStyle_toc = OxmlElement('w:pStyle')
    pStyle_toc.set(f'{{{nsmap_w}}}val', 'TOC1')
    pPr_toc.append(pStyle_toc)
    p_toc.append(pPr_toc)
    
    # fldChar begin
    r_begin = OxmlElement('w:r')
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(f'{{{nsmap_w}}}fldCharType', 'begin')
    r_begin.append(fc_begin)
    p_toc.append(r_begin)
    
    # instrText
    r_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_instr.append(instr)
    p_toc.append(r_instr)
    
    # fldChar separate
    r_sep = OxmlElement('w:r')
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(f'{{{nsmap_w}}}fldCharType', 'separate')
    r_sep.append(fc_sep)
    p_toc.append(r_sep)
    
    # Placeholder text (提示用户更新域)
    r_hint = OxmlElement('w:r')
    rPr_hint = OxmlElement('w:rPr')
    rFonts_hint = OxmlElement('w:rFonts')
    rFonts_hint.set(f'{{{nsmap_w}}}ascii', 'Times New Roman')
    rFonts_hint.set(f'{{{nsmap_w}}}hAnsi', 'Times New Roman')
    rPr_hint.append(rFonts_hint)
    r_hint.append(rPr_hint)
    t_hint = OxmlElement('w:t')
    t_hint.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_hint.text = '（请在 Word 中右键此处 → 更新域，自动生成目录）'
    r_hint.append(t_hint)
    p_toc.append(r_hint)
    
    # fldChar end
    r_end = OxmlElement('w:r')
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(f'{{{nsmap_w}}}fldCharType', 'end')
    r_end.append(fc_end)
    p_toc.append(r_end)
    
    new_content.append(p_toc)
    
    # 替换 SDT 的 sdtContent
    old_content = sdt_to_replace.find(f'{{{nsmap_w}}}sdtContent')
    if old_content is not None:
        sdt_to_replace.replace(old_content, new_content)
    else:
        sdt_to_replace.append(new_content)
    
    print("✅ 目录 (TOC) 已更新为新的域代码，请在 Word 中右键 → 更新域")

# ============================================================
# 保存
# ============================================================

doc.save(OUTPUT)
print(f"✅ 文档已保存至: {OUTPUT}")
